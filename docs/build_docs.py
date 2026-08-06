# -*- coding: utf-8 -*-
"""
Render the conversation transcript to a colour-designed PDF and DOCX.

Usage:
    py docs/build_docs.py

Outputs (next to this script):
    Multimodal_RAG_Design_Conversation.pdf
    Multimodal_RAG_Design_Conversation.docx
"""

import os
import re
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

import importlib  # noqa: E402
from pathlib import Path  # noqa: E402

# What to render. Either a python content module, or a Markdown file.
#   py docs/build_docs.py                        -> conversation transcript
#   py docs/build_docs.py asbuilt_content        -> as-built architecture
#   py docs/build_docs.py CODE_WALKTHROUGH.md    -> rendered from Markdown
_ARG = sys.argv[1] if len(sys.argv) > 1 and not sys.argv[1].startswith("-") \
    else "conversation_content"

if _ARG.lower().endswith(".md"):
    from md_to_blocks import build as _build_from_md  # noqa: E402

    _path = Path(_ARG)
    if not _path.is_absolute():
        _path = Path(HERE) / _path
    BLOCKS, TITLE = _build_from_md(_path)
    _BASENAME = _path.stem
else:
    _content = importlib.import_module(_ARG)
    BLOCKS = _content.BLOCKS
    _BASENAME = getattr(_content, "OUTPUT_BASENAME", "Multimodal_RAG_Design_Conversation")
    TITLE = getattr(_content, "DOC_TITLE", "Multimodal RAG Pipeline - Design Conversation")

OUT_PDF = os.path.join(HERE, _BASENAME + ".pdf")
OUT_DOCX = os.path.join(HERE, _BASENAME + ".docx")

# ----------------------------------------------------------------------
# Shared design system
# ----------------------------------------------------------------------
INDIGO = "4F46E5"
VIOLET = "7C3AED"
TEAL = "0D9488"
AMBER = "D97706"
ROSE = "E11D48"
CYAN = "0891B2"
SLATE = "1E293B"
GREY = "64748B"
WHITE = "FFFFFF"

BG_CODE = "F1F5F9"
BG_DIAGRAM = "F8FAFF"
BG_TABLE_ALT = "F5F7FF"
BG_USER = "EEF2FF"
BG_ASST = "ECFDF5"
BG_SYS = "FFF7ED"

TURN_STYLE = {
    "user": (BG_USER, INDIGO, "USER"),
    "assistant": (BG_ASST, TEAL, "RESPONSE"),
    "system": (BG_SYS, AMBER, "DECISION"),
}

CALLOUT_STYLE = {
    "info": ("ECFEFF", CYAN),
    "key": ("EEF2FF", INDIGO),
    "warn": ("FFFBEB", AMBER),
    "tip": ("F0FDFA", TEAL),
}

# Characters that some fonts lack -> safe equivalents
SANITIZE = {
    "\u26a0": "!", "\u21d2": "=>", "\u2713": "OK", "\u2717": "x",
    "\u2192": "->", "\u2190": "<-", "\u2265": ">=", "\u2264": "<=",
}


def clean(text):
    for bad, good in SANITIZE.items():
        text = text.replace(bad, good)
    return text


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)", re.S)


def parse_inline(text):
    """Split text into [(chunk, style)] where style is '', 'b' or 'code'."""
    out = []
    for part in INLINE_RE.split(clean(text)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], "b"))
        elif part.startswith("`") and part.endswith("`"):
            out.append((part[1:-1], "code"))
        else:
            out.append((part, ""))
    return out


def col_weights(headers, rows, min_share=0.09):
    """Proportional column widths based on average content length."""
    n = len(headers)
    lengths = []
    for i in range(n):
        cells = [str(headers[i])] + [str(r[i]) for r in rows if i < len(r)]
        avg = sum(len(c) for c in cells) / max(len(cells), 1)
        lengths.append(max(avg, 4))
    total = sum(lengths)
    shares = [max(l / total, min_share) for l in lengths]
    s = sum(shares)
    return [x / s for x in shares]


# ======================================================================
# PDF renderer
# ======================================================================
def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table,
        TableStyle, Preformatted, PageBreak, KeepTogether, HRFlowable,
        ListFlowable, ListItem,
    )

    F = r"C:\Windows\Fonts"

    def reg(name, filename, fallback=None):
        path = os.path.join(F, filename)
        if not os.path.exists(path) and fallback:
            path = os.path.join(F, fallback)
        pdfmetrics.registerFont(TTFont(name, path))

    reg("Body", "calibri.ttf")
    reg("Body-B", "calibrib.ttf", "calibri.ttf")
    reg("Body-I", "calibrii.ttf", "calibri.ttf")
    reg("Body-BI", "calibriz.ttf", "calibrib.ttf")
    reg("Mono", "consola.ttf")
    reg("Mono-B", "consolab.ttf", "consola.ttf")
    pdfmetrics.registerFontFamily("Body", normal="Body", bold="Body-B",
                                  italic="Body-I", boldItalic="Body-BI")
    pdfmetrics.registerFontFamily("Mono", normal="Mono", bold="Mono-B",
                                  italic="Mono", boldItalic="Mono-B")

    C = lambda h: colors.HexColor("#" + h)

    PW, PH = A4
    ML = MR = 17 * mm
    MT = 16 * mm
    MB = 16 * mm
    USABLE = PW - ML - MR

    def esc(t):
        return (t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def rich(text):
        """Inline markup -> reportlab mini-HTML."""
        out = []
        for chunk, style in parse_inline(text):
            e = esc(chunk)
            if style == "b":
                out.append("<b>%s</b>" % e)
            elif style == "code":
                out.append('<font face="Mono" size="8.6" color="#%s">%s</font>' % (ROSE, e))
            else:
                out.append(e)
        return "".join(out)

    S = {}
    S["p"] = ParagraphStyle("p", fontName="Body", fontSize=10.2, leading=15.2,
                            textColor=C(SLATE), spaceAfter=7, alignment=TA_LEFT)
    S["h1"] = ParagraphStyle("h1", fontName="Body-B", fontSize=19, leading=23,
                             textColor=C(INDIGO), spaceBefore=6, spaceAfter=3)
    S["h2"] = ParagraphStyle("h2", fontName="Body-B", fontSize=14.5, leading=18,
                             textColor=C(TEAL), spaceBefore=14, spaceAfter=5)
    S["h3"] = ParagraphStyle("h3", fontName="Body-B", fontSize=11.8, leading=15,
                             textColor=C(AMBER), spaceBefore=11, spaceAfter=4)
    S["li"] = ParagraphStyle("li", parent=S["p"], spaceAfter=4)
    S["mono"] = ParagraphStyle("mono", fontName="Mono", fontSize=8.2, leading=11.2,
                               textColor=C(SLATE))
    S["diag"] = ParagraphStyle("diag", fontName="Mono", fontSize=7.4, leading=8.9,
                               textColor=C(SLATE))
    S["cap"] = ParagraphStyle("cap", fontName="Body-BI", fontSize=9, leading=12,
                              textColor=C(VIOLET), spaceAfter=3)
    S["th"] = ParagraphStyle("th", fontName="Body-B", fontSize=9.2, leading=12,
                             textColor=colors.white)
    S["td"] = ParagraphStyle("td", fontName="Body", fontSize=9.0, leading=12.4,
                             textColor=C(SLATE))
    S["cotitle"] = ParagraphStyle("cotitle", fontName="Body-B", fontSize=10.4,
                                  leading=13.6, textColor=C(SLATE), spaceAfter=2)
    S["cobody"] = ParagraphStyle("cobody", fontName="Body", fontSize=9.8,
                                 leading=14, textColor=C(SLATE))
    S["turnlabel"] = ParagraphStyle("turnlabel", fontName="Body-B", fontSize=13.5,
                                    leading=17, textColor=C(SLATE))
    S["turnsub"] = ParagraphStyle("turnsub", fontName="Body-I", fontSize=9.6,
                                  leading=13, textColor=C(GREY))
    S["turntag"] = ParagraphStyle("turntag", fontName="Body-B", fontSize=8.2,
                                  leading=11, textColor=colors.white)
    S["covert"] = ParagraphStyle("covert", fontName="Body-B", fontSize=31,
                                 leading=36, textColor=colors.white)
    S["covers"] = ParagraphStyle("covers", fontName="Body", fontSize=13.5,
                                 leading=18, textColor=colors.white)
    S["metak"] = ParagraphStyle("metak", fontName="Body-B", fontSize=9.4,
                                leading=13, textColor=C(INDIGO))
    S["metav"] = ParagraphStyle("metav", fontName="Body", fontSize=9.4,
                                leading=13, textColor=C(SLATE))

    story = []

    def block_wrap(inner, bg, accent, pad=9, accent_w=3.4):
        t = Table([[inner]], colWidths=[USABLE])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), C(bg)),
            ("LINEBEFORE", (0, 0), (0, -1), accent_w, C(accent)),
            ("LEFTPADDING", (0, 0), (-1, -1), pad + 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), pad),
            ("TOPPADDING", (0, 0), (-1, -1), pad),
            ("BOTTOMPADDING", (0, 0), (-1, -1), pad),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    def add_table(headers, rows):
        w = [x * USABLE for x in col_weights(headers, rows)]
        data = [[Paragraph(rich(str(h)), S["th"]) for h in headers]]
        for r in rows:
            data.append([Paragraph(rich(str(c)), S["td"]) for c in r])
        t = Table(data, colWidths=w, repeatRows=1)
        st = [
            ("BACKGROUND", (0, 0), (-1, 0), C(INDIGO)),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.4, C("D8DEF0")),
            ("LINEBELOW", (0, 0), (-1, 0), 1.1, C(VIOLET)),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]
        for i in range(1, len(data)):
            if i % 2 == 0:
                st.append(("BACKGROUND", (0, i), (-1, i), C(BG_TABLE_ALT)))
        t.setStyle(TableStyle(st))
        story.append(Spacer(1, 3))
        story.append(t)
        story.append(Spacer(1, 9))

    for kind, payload in BLOCKS:

        if kind == "cover":
            story.append(Spacer(1, 24 * mm))
            band = Table([[Paragraph(clean(payload["title"]), S["covert"])],
                          [Paragraph(clean(payload["subtitle"]), S["covers"])]],
                         colWidths=[USABLE])
            band.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), C(INDIGO)),
                ("LINEBEFORE", (0, 0), (0, -1), 8, C(VIOLET)),
                ("LEFTPADDING", (0, 0), (-1, -1), 22),
                ("RIGHTPADDING", (0, 0), (-1, -1), 18),
                ("TOPPADDING", (0, 0), (0, 0), 26),
                ("BOTTOMPADDING", (0, 1), (0, 1), 26),
            ]))
            story.append(band)
            strip = Table([[""] * 4], colWidths=[USABLE / 4.0] * 4, rowHeights=[7])
            strip.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, 0), C(VIOLET)),
                ("BACKGROUND", (1, 0), (1, 0), C(CYAN)),
                ("BACKGROUND", (2, 0), (2, 0), C(TEAL)),
                ("BACKGROUND", (3, 0), (3, 0), C(AMBER)),
            ]))
            story.append(strip)
            story.append(Spacer(1, 16))
            meta = [[Paragraph(k, S["metak"]), Paragraph(clean(v), S["metav"])]
                    for k, v in payload["meta"]]
            mt = Table(meta, colWidths=[USABLE * 0.22, USABLE * 0.78])
            mt.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LINEBELOW", (0, 0), (-1, -2), 0.4, C("E2E8F0")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ]))
            story.append(mt)
            story.append(PageBreak())

        elif kind == "turn":
            role, label, sub = payload
            bg, accent, tag = TURN_STYLE[role]
            tagcell = Table([[Paragraph(tag, S["turntag"])]], colWidths=[62])
            tagcell.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), C(accent)),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            inner = Table([[tagcell, Paragraph(clean(label), S["turnlabel"])],
                           ["", Paragraph(clean(sub), S["turnsub"])]],
                          colWidths=[70, USABLE - 70 - 26])
            inner.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]))
            story.append(Spacer(1, 10))
            story.append(block_wrap(inner, bg, accent, pad=10, accent_w=5))
            story.append(Spacer(1, 10))

        elif kind == "h1":
            story.append(Spacer(1, 8))
            story.append(Paragraph(rich(payload), S["h1"]))
            story.append(HRFlowable(width="100%", thickness=2.2, color=C(VIOLET),
                                    spaceBefore=2, spaceAfter=9))

        elif kind == "h2":
            story.append(Paragraph(rich(payload), S["h2"]))

        elif kind == "h3":
            story.append(Paragraph(rich(payload), S["h3"]))

        elif kind == "p":
            story.append(Paragraph(rich(payload), S["p"]))

        elif kind in ("bullets", "numbers"):
            items = [ListItem(Paragraph(rich(x), S["li"]), leftIndent=14)
                     for x in payload]
            story.append(ListFlowable(
                items,
                bulletType="bullet" if kind == "bullets" else "1",
                start="circle" if kind == "bullets" else None,
                bulletColor=C(VIOLET), bulletFontSize=8.5,
                leftIndent=15, bulletOffsetY=0.5,
            ))
            story.append(Spacer(1, 7))

        elif kind == "code":
            body = Preformatted(clean(payload), S["mono"])
            story.append(KeepTogether(block_wrap(body, BG_CODE, CYAN)))
            story.append(Spacer(1, 9))

        elif kind == "diagram":
            title, body = payload
            grp = [Paragraph(clean(title), S["cap"]),
                   block_wrap(Preformatted(clean(body), S["diag"]),
                              BG_DIAGRAM, VIOLET)]
            story.append(KeepTogether(grp))
            story.append(Spacer(1, 10))

        elif kind == "table":
            add_table(*payload)

        elif kind == "kv":
            rows = [[Paragraph(rich(k), S["metak"]), Paragraph(rich(v), S["td"])]
                    for k, v in payload]
            t = Table(rows, colWidths=[USABLE * 0.20, USABLE * 0.80])
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (0, -1), C("F8FAFC")),
                ("LINEBELOW", (0, 0), (-1, -2), 0.4, C("E2E8F0")),
                ("LINEBEFORE", (0, 0), (0, -1), 2.4, C(TEAL)),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(Spacer(1, 2))
            story.append(t)
            story.append(Spacer(1, 9))

        elif kind == "callout":
            ck, ctitle, cbody = payload
            bg, accent = CALLOUT_STYLE[ck]
            inner = Table([[Paragraph(rich(ctitle), S["cotitle"])],
                           [Paragraph(rich(cbody), S["cobody"])]],
                          colWidths=[USABLE - 26])
            inner.setStyle(TableStyle([
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(KeepTogether(block_wrap(inner, bg, accent, pad=9, accent_w=4)))
            story.append(Spacer(1, 10))

        elif kind == "rule":
            story.append(HRFlowable(width="100%", thickness=0.8, color=C("CBD5E1"),
                                    spaceBefore=6, spaceAfter=8))

        elif kind == "pagebreak":
            story.append(PageBreak())

    def decorate(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(C(INDIGO))
        canvas.rect(0, PH - 5 * mm, PW, 5 * mm, stroke=0, fill=1)
        canvas.setFillColor(C(VIOLET))
        canvas.rect(0, PH - 5 * mm, PW * 0.34, 5 * mm, stroke=0, fill=1)
        canvas.setFillColor(C(AMBER))
        canvas.rect(0, PH - 5 * mm, PW * 0.12, 5 * mm, stroke=0, fill=1)
        if doc.page > 1:
            canvas.setFont("Body", 8)
            canvas.setFillColor(C(GREY))
            canvas.drawString(ML, 9 * mm, TITLE)
            canvas.drawRightString(PW - MR, 9 * mm, "Page %d" % doc.page)
            canvas.setStrokeColor(C("E2E8F0"))
            canvas.setLineWidth(0.5)
            canvas.line(ML, 12 * mm, PW - MR, 12 * mm)
        canvas.restoreState()

    doc = BaseDocTemplate(OUT_PDF, pagesize=A4,
                          leftMargin=ML, rightMargin=MR,
                          topMargin=MT, bottomMargin=MB,
                          title=TITLE, author="Claude Code")
    frame = Frame(ML, MB, USABLE, PH - MT - MB, id="body",
                  leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=decorate)])
    doc.build(story)
    return OUT_PDF


# ======================================================================
# DOCX renderer
# ======================================================================
def build_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def rgb(h):
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def shade(element, hexcolor):
        pr = element.get_or_add_tcPr() if element.tag.endswith("}tc") else \
            element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hexcolor)
        pr.append(shd)

    def para_shade(p, hexcolor):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hexcolor)
        p._p.get_or_add_pPr().append(shd)

    def para_border(p, color, sides=("left",), size=18, space=6):
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        for s in sides:
            el = OxmlElement("w:" + s)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), str(space))
            el.set(qn("w:color"), color)
            pbdr.append(el)
        pPr.append(pbdr)

    def cell_shade(cell, hexcolor):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hexcolor)
        cell._tc.get_or_add_tcPr().append(shd)

    def cell_border(cell, color, sides=("left",), size=24):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for s in sides:
            el = OxmlElement("w:" + s)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            borders.append(el)
        tcPr.append(borders)

    def fixed_layout(t):
        """Stop Word from auto-fitting; honour the widths we set."""
        tblPr = t._tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)

    def no_space(p, before=0, after=0, line=None):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if line:
            pf.line_spacing = line

    def write_runs(p, text, base_size=10.5, base_color=SLATE,
                   base_font="Calibri", bold_all=False):
        for chunk, style in parse_inline(text):
            r = p.add_run(chunk)
            if style == "code":
                r.font.name = "Consolas"
                r.font.size = Pt(base_size - 1.2)
                r.font.color.rgb = rgb(ROSE)
            else:
                r.font.name = base_font
                r.font.size = Pt(base_size)
                r.font.color.rgb = rgb(base_color)
                r.bold = bold_all or (style == "b")
        return p

    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(1.9)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    CONTENT_W = Cm(21.0 - 3.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(SLATE)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # footer
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(TITLE)
    fr.font.size = Pt(8)
    fr.font.color.rgb = rgb(GREY)

    def banner(text, size, color, bg=None, before=6, after=6, border=None):
        p = doc.add_paragraph()
        no_space(p, before, after, 1.0)
        if bg:
            para_shade(p, bg)
        if border:
            para_border(p, border, ("left",), 24, 8)
        write_runs(p, text, base_size=size, base_color=color, bold_all=True)
        return p

    def mono_block(text, bg, accent, size=7.6):
        for i, line in enumerate(clean(text).split("\n")):
            p = doc.add_paragraph()
            no_space(p, 0, 0, 1.0)
            para_shade(p, bg)
            para_border(p, accent, ("left",), 18, 8)
            p.paragraph_format.left_indent = Cm(0.25)
            r = p.add_run(line if line else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(size)
            r.font.color.rgb = rgb(SLATE)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            if i == 0:
                p.paragraph_format.space_before = Pt(5)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(9)

    def make_table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.autofit = False
        weights = col_weights(headers, rows)
        widths = [Cm(w * (21.0 - 3.8)) for w in weights]

        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            cell_shade(hdr[i], INDIGO)
            p = hdr[i].paragraphs[0]
            no_space(p, 2, 2, 1.0)
            write_runs(p, str(h), base_size=9.5, base_color=WHITE, bold_all=True)

        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci, val in enumerate(row):
                if ri % 2 == 1:
                    cell_shade(cells[ci], BG_TABLE_ALT)
                p = cells[ci].paragraphs[0]
                no_space(p, 2, 2, 1.05)
                write_runs(p, str(val), base_size=9.2)

        t.style = "Table Grid"
        for row in t.rows:
            for i, c in enumerate(row.cells):
                c.width = widths[i]
        fixed_layout(t)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        return t

    def one_cell(bg, accent, builder, accent_size=30):
        t = doc.add_table(rows=1, cols=1)
        t.autofit = False
        cell = t.rows[0].cells[0]
        cell.width = CONTENT_W
        cell_shade(cell, bg)
        cell_border(cell, accent, ("left",), accent_size)
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        builder(cell)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        return t

    for kind, payload in BLOCKS:

        if kind == "cover":
            for _ in range(2):
                doc.add_paragraph().paragraph_format.space_after = Pt(0)

            def cover_cell(cell):
                p = cell.add_paragraph()
                no_space(p, 14, 4, 1.0)
                write_runs(p, payload["title"], base_size=30, base_color=WHITE, bold_all=True)
                p2 = cell.add_paragraph()
                no_space(p2, 0, 14, 1.0)
                write_runs(p2, payload["subtitle"], base_size=13.5, base_color=WHITE)

            one_cell(INDIGO, VIOLET, cover_cell, accent_size=48)

            strip = doc.add_table(rows=1, cols=4)
            strip.autofit = False
            for i, col in enumerate([VIOLET, CYAN, TEAL, AMBER]):
                c = strip.rows[0].cells[i]
                c.width = Cm((21.0 - 3.8) / 4)
                cell_shade(c, col)
                pp = c.paragraphs[0]
                no_space(pp, 0, 0, 1.0)
                pp.add_run(" ").font.size = Pt(3)
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

            mt = doc.add_table(rows=0, cols=2)
            mt.autofit = False
            for k, v in payload["meta"]:
                cells = mt.add_row().cells
                cells[0].width = Cm((21.0 - 3.8) * 0.24)
                cells[1].width = Cm((21.0 - 3.8) * 0.76)
                p = cells[0].paragraphs[0]
                no_space(p, 3, 3, 1.0)
                write_runs(p, k, base_size=9.5, base_color=INDIGO, bold_all=True)
                p = cells[1].paragraphs[0]
                no_space(p, 3, 3, 1.0)
                write_runs(p, v, base_size=9.5)
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        elif kind == "turn":
            role, label, sub = payload
            bg, accent, tag = TURN_STYLE[role]

            def turn_cell(cell, label=label, sub=sub, tag=tag, accent=accent):
                p = cell.add_paragraph()
                no_space(p, 6, 1, 1.0)
                r = p.add_run(tag + "  ")
                r.font.name = "Calibri"
                r.font.size = Pt(8.5)
                r.bold = True
                r.font.color.rgb = rgb(accent)
                write_runs(p, label, base_size=13.5, bold_all=True)
                p2 = cell.add_paragraph()
                no_space(p2, 0, 6, 1.0)
                write_runs(p2, sub, base_size=9.5, base_color=GREY)
                p2.runs[0].italic = True

            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            one_cell(bg, accent, turn_cell, accent_size=40)

        elif kind == "h1":
            p = banner(payload, 19, INDIGO, before=12, after=2)
            para_border(p, VIOLET, ("bottom",), 18, 4)

        elif kind == "h2":
            banner(payload, 14.5, TEAL, before=13, after=4)

        elif kind == "h3":
            banner(payload, 11.8, AMBER, before=10, after=3)

        elif kind == "p":
            p = doc.add_paragraph()
            write_runs(p, payload)

        elif kind in ("bullets", "numbers"):
            style = "List Bullet" if kind == "bullets" else "List Number"
            for item in payload:
                p = doc.add_paragraph(style=style)
                no_space(p, 1, 3, 1.15)
                write_runs(p, item, base_size=10.3)

        elif kind == "code":
            mono_block(payload, BG_CODE, CYAN, size=8.4)

        elif kind == "diagram":
            title, body = payload
            p = doc.add_paragraph()
            no_space(p, 8, 2, 1.0)
            r = p.add_run(clean(title))
            r.font.name = "Calibri"
            r.font.size = Pt(9.5)
            r.bold = True
            r.italic = True
            r.font.color.rgb = rgb(VIOLET)
            mono_block(body, BG_DIAGRAM, VIOLET, size=7.4)

        elif kind == "table":
            make_table(*payload)

        elif kind == "kv":
            t = doc.add_table(rows=0, cols=2)
            t.autofit = False
            for k, v in payload:
                cells = t.add_row().cells
                cells[0].width = Cm((21.0 - 3.8) * 0.22)
                cells[1].width = Cm((21.0 - 3.8) * 0.78)
                cell_shade(cells[0], "F8FAFC")
                cell_border(cells[0], TEAL, ("left",), 18)
                p = cells[0].paragraphs[0]
                no_space(p, 3, 3, 1.05)
                write_runs(p, k, base_size=9.5, base_color=INDIGO, bold_all=True)
                p = cells[1].paragraphs[0]
                no_space(p, 3, 3, 1.1)
                write_runs(p, v, base_size=9.4)
            t.style = "Table Grid"
            fixed_layout(t)
            doc.add_paragraph().paragraph_format.space_after = Pt(3)

        elif kind == "callout":
            ck, ctitle, cbody = payload
            bg, accent = CALLOUT_STYLE[ck]

            def co_cell(cell, ctitle=ctitle, cbody=cbody):
                p = cell.add_paragraph()
                no_space(p, 5, 2, 1.05)
                write_runs(p, ctitle, base_size=10.5, bold_all=True)
                p2 = cell.add_paragraph()
                no_space(p2, 0, 5, 1.15)
                write_runs(p2, cbody, base_size=9.9)

            one_cell(bg, accent, co_cell, accent_size=34)

        elif kind == "rule":
            p = doc.add_paragraph()
            no_space(p, 4, 6)
            para_border(p, "CBD5E1", ("bottom",), 6, 1)

        elif kind == "pagebreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    pdf = build_pdf()
    print("PDF  ->", pdf, "%.1f KB" % (os.path.getsize(pdf) / 1024))
    docx = build_docx()
    print("DOCX ->", docx, "%.1f KB" % (os.path.getsize(docx) / 1024))
